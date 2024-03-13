# -*- coding: utf-8 -*-
"""
Meraki Cat webexteamsbot is a chat bot with a swiss army knife
of functions for checking & translating Catalyst IOSXE to Meraki
switch configs, registering Catalyst switches to Dashboard and
claiming them.
"""
import ngrok
import meraki
from ciscoconfparse2 import CiscoConfParse
from webexteamsbot import TeamsBot
from webexteamsbot.models import Response
from netmiko import ConnectHandler
from pyadaptivecards.card import AdaptiveCard
from pyadaptivecards.components import TextBlock, Column
from pyadaptivecards.container import ColumnSet, Container
from pyadaptivecards.options import Colors, FontSize, HorizontalAlignment
from docx2pdf import convert
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict
from functools import reduce
from mc_cfg_check import CheckFeatures
from mc_translate import Evaluate, Meraki_config_down#, Meraki_config_up
from mc_claim import Claim
from mc_register import Register
from mc_splitcheck_serials import Split_check_serials
from mc_ping import Ping
from mc_file_exists import File_exists
from mc_user_info import *
from mc_cards import *

from tabulate import tabulate
tabulate.PRESERVE_WHITESPACE = True
import os, platform, requests, json, pprint, re, sys, time

# Retrieve required Meraki details from environment variables
meraki_api_key = os.getenv("MERAKI_API_KEY")
# If the required details were not in the environment variables
# grab them from the mc_user_info.py file
if meraki_api_key is None:
    meraki_api_key = MERAKI_API_KEY

# Retrieve required SSH details from environment variables
ios_username = os.getenv("IOS_USERNAME")
ios_password = os.getenv("IOS_PASSWORD")
ios_secret = os.getenv("IOS_SECRET")
ios_port = os.getenv("IOS_PORT")
# If the required details were not in the environment variables
# grab them from the mc_user_info.py file
if ios_username is None:
    ios_username = IOS_USERNAME
if ios_password is None:
    ios_password = IOS_PASSWORD
if ios_secret is None:
    ios_secret = IOS_SECRET
if ios_port is None:
    ios_port = IOS_PORT
    if ios_port is None:
        ios_port = 22

# If we were run without arguments, run as a BOT
# Otherwise, we will attempt to use the args in batch mode
BOT = False
if len(sys.argv) == 1:
    BOT = True

teams_emails = list()
if BOT:
    # Retrieve required details from environment variables
    bot_email = os.getenv("TEAMS_BOT_EMAIL")
    bot_app_name = os.getenv("TEAMS_BOT_APP_NAME")
    teams_token = os.getenv("TEAMS_BOT_TOKEN")
    if not os.getenv("TEAMS_EMAILS") is None:
        teams_emails.append(os.getenv("TEAMS_EMAILS"))
    ngrok_token = os.getenv("NGROK_AUTHTOKEN")
    print(f"ngrok_token = {ngrok_token}")
    #bot_url = os.getenv("TEAMS_BOT_URL")
    
    # If the required details were not in the environment variables
    # grab them from the mc_user_info.py file
    if bot_email is None:
        bot_email = TEAMS_BOT_EMAIL
    if bot_app_name is None:
        bot_app_name = TEAMS_BOT_APP_NAME
    if teams_token is None:
        teams_token = TEAMS_BOT_TOKEN
    if len(teams_emails) == 0:
        teams_emails = TEAMS_EMAILS
    if ngrok_token == None:
        ngrok_token = NGROK_AUTHTOKEN
    listener = ngrok.forward("localhost:5000", authtoken=ngrok_token)
    print(f"Ingress established at: {listener.url()}");
    bot_url = listener.url()
    #if bot_url is None:
    #    bot_url = BOT_URL

    # Either way, let's got the Bot's first name in case we are
    # directly addressed in room with multiple users
    bot_fname = bot_app_name.split()[0].strip()

debug = DEBUG or DEBUG_MAIN

# Setup some global variables
payload = {}
organizations = {}
api = ""
payload = None
configured_ports = defaultdict(list)
unconfigured_ports = defaultdict(list)
command_line_msg = Response()
times = False
report = False

#Setup some global, stateful variables
config_file = ""
host_id = ""
nm_list = list()
meraki_serials = list()
meraki_orgs = list()
meraki_networks = list()
meraki_org=""
meraki_org_name=""
meraki_net=""
meraki_net_name=""
meraki_urls = list()

# Request the lists of Organizations and their Networks from Dashboard
if not meraki_api_key == "":
    if debug:
        print(f"Trying to setup a dashboard instance")
    try:
        dashboard = meraki.DashboardAPI(api_key=meraki_api_key,output_log=False,suppress_logging=True)
    except:
        print("We were unable to login to the dashboard.")
        sys.exit()
    if debug:
        print(f"Got one, now trying to get the list of Orgs")
    # Even though, right now this app only supports a single Org...
    try:
        meraki_orgs = dashboard.organizations.getOrganizations()
    except:
        print("We were unable to get the list of Orgs.")
        sys.exit()
    if debug:
        print(f"meraki_orgs = {meraki_orgs}")
    x = 0
    while x <= len(meraki_orgs)-1:
        try:
            raw_nets = dashboard.organizations.getOrganizationNetworks(organizationId=meraki_orgs[x]['id'])
        except:
            print(f"We were unable to get the list of networks for {meraki_orgs[x]['name']}.")
            sys.exit()
        if debug:
            print(raw_nets)
        y = 0
        while y <= len(raw_nets)-1:
            meraki_networks.append(raw_nets[y])
            y+=1
        x+=1
    if debug:
        print(f"meraki_networks = {meraki_networks}")
    if len(meraki_orgs) == 1:
        meraki_org = meraki_orgs[0]['id']
        meraki_org_name = meraki_orgs[0]['name']
        if debug:
            print(f"meraki_org = {meraki_org}")
            print(f"meraki_org_name = {meraki_org_name}")
        if len(meraki_orgs) == 1 and len(meraki_networks) == 1:
            meraki_net = meraki_networks[0]['id']
            meraki_net_name = meraki_networks[0]['name']
            if debug:
                print(f"meraki_net = {meraki_net}")
                print(f"meraki_net_name = {meraki_net_name}")


if BOT:
    # If any of the required bot variables are missing, terminate the app
    if not bot_email or not teams_token or not bot_url or not bot_app_name:
        print(
            "merakicat.py - Missing Environment Variable. Please see the 'Usage'"
            " section in the README."
        )
        if not bot_email:
            print("TEAMS_BOT_EMAIL")
        if not teams_token:
            print("TEAMS_BOT_TOKEN")
        if not bot_app_name:
            print("TEAMS_BOT_APP_NAME")
        if not bot_url:
            print("TEAMS_BOT_URL")
        sys.exit()
    
    
    # Create a Bot Object
    #   Note: debug mode prints out more details about processing to terminal
    #   Note: the `approved_users=approved_users` line commented out and shown as reference
    #
    # Example: How to limit the approved Webex Teams accounts for interaction
    # List of email accounts of approved users to talk with the bot
    # approved_users = [
    #     "josmith@demo.local",
    # ]
    
    if debug:
        print(f"teams_emails = {teams_emails}")
    bot = TeamsBot(
        bot_app_name,
        teams_bot_token=teams_token,
        teams_bot_url=bot_url,
        teams_bot_email=bot_email,
        debug=True,
        # Comment out the approved_users lines if you don't care...
        approved_users = teams_emails,
        webhook_resource_event=[
            {"resource": "messages", "event": "created"},
            {"resource": "attachmentActions", "event": "created"},
        ],
    )


# Create a custom bot greeting function, returned when no /command is given.
# The default behavior of the bot is to return the '/help' command response
# If there is an English language command line, try to work with that.

def greeting(incoming_msg):
    global config_file, host_id, meraki_net, meraki_net_name, meraki_serials
    global times, report
    
    if debug:
        print(f"incoming_msg = {incoming_msg}")
        print(f"type of incoming_msg = {type(incoming_msg)}")
    
    # Create a Response object to later craft a reply in Markdown.
    response = Response()
    
    # This will be our copy of the user input to work with
    user_text = incoming_msg.text
    
    # If first word from the user's input was Bot's first name, remove it
    if BOT:
        if user_text.split()[0] == bot_fname:
            user_text = user_text.split(bot_fname + ' ', 1)[1]
    
    # Grab the first word from the user's input.
    command = user_text.split()[0].lower()
    
    print(f"command = {command}")
    
    # If the user asked for a report, we will try to give it to them
    report = False
    if not re.search('\swith\sreports|\swith\sreport|\swith\sreporting', user_text, re.IGNORECASE) == None:
        user_text = re.sub('\swith\sreports|\swith\sreports|\swith\sreporting', '', user_text, re.IGNORECASE)
        report = True
    
    # If the user asked for timing, we will try to give it to them
    times = False
    if not re.search('\swith\stimings|\swith\stiming|\swith\stime|\swith\stimes', user_text, re.IGNORECASE) == None:
        user_text = re.sub('\swith\stimings|\swith\stiming|\swith\stimes|\swith\stime', '', user_text, re.IGNORECASE)
        times = True
    
    serials = list()
    
    # Test if it is equivalent to a /command.
    match command: 
        case "migrate":
          # If the only thing the user typed was "migrate""...
          if user_text.lower() == 'migrate':
              # Check and see if we have a global stateful host and network to work with
              if host_id=="" and meraki_net=="":
                  # We did not...
                  response.markdown = "I'm sorry, but I don't have a host or a Network that we are working with."
              elif host_id=="":
                  # Just missing the host
                  response.markdown = "I'm sorry, but I don't have a host that we are working with."
              elif meraki_net=="":
                  # Just missing the Network
                  response.markdown = "I'm sorry, but I don't have a Network that we are working with."
              else:
                  # We did, so migrate it!
                  response.markdown = migrate_switch(incoming_msg,host=host_id,dest_net=meraki_net)
          # Well, did they type more after "migrate" ?
          elif user_text.lower().startswith('migrate'):
              dest_net = meraki_net
              # Did they enter a network name after "to"?
              if re.search('to ', user_text, re.IGNORECASE):
                  regex = re.compile(r"\s*to\s *", flags=re.I)
                  if not regex.split(user_text)[1] == "":
                      # They did, so register it!
                      dest_net_name = regex.split(user_text)[1]
                      dest_net = ""
                      test = [d['id'] for d in meraki_networks if d['name']==dest_net_name]
                      if not len(test) == 0:
                          dest_net = test[0]
                          meraki_net_name = dest_net_name
                      if dest_net == "":
                          response.markdown = "I'm sorry, but {} is not in your list of Meraki networks.".format(dest_net_name)
                      else:
                          if debug:
                              print(f"dest_net = {dest_net}")
              if dest_net == "":
                  response.markdown = "You need to enter a Meraki network to register into."
              host = host_id
              # Did they type "translate host <something>" ?
              if re.search('host ', user_text, re.IGNORECASE):
                  if debug:
                      print(f"I made it to host...")
                  if not user_text.split("host ",1)[1] == "":
                      # They did, so translate it!
                      maybe_host = user_text.split("host ",1)[1]
                      regex = re.compile(r"\s*to\s *", flags=re.I)
                      the_rest = regex.split(maybe_host)[0]
                      host = the_rest.strip()
                      if debug:
                          print(f"host = {host}")
                      if not Ping(host):
                          response.markdown = "I was unable to ping that host."
                      response.markdown = migrate_switch(incoming_msg,host=host,dest_net=dest_net)
                  else:
                      # They did not, so BUMP the user.
                      response.markdown = "I'm sorry, but I don't have a host that we are working with.  Use the **/check** command."
              if host =="":
                      response.markdown = "I'm sorry, but I don't have a host that we are working with.  Use the **/check** command."
          else:
              response.markdown = migrate_switch(incoming_msg,dest_net=dest_net)
        
        case "translate":
          # If the only thing the user typed was "translate""...
          if user_text.lower() == 'translate':
              # Check and see if we have a global stateful config filespec to work with
              if config_file=="" and host_id=="":
                  # We did not...
                  response.markdown = "I'm sorry, but I don't have a config or host that we are working with.  Use the **/check** command."
              else:
                  if not len(meraki_serials) == 0:
                      # We did, so translate it!
                      serials = meraki_serials
                      response.markdown = translate_switch(incoming_msg,config=config_file,host=host_id,serials=serials)
                  else:
                      response.markdown = "I'm sorry, but I don't have a list of Meraki switch serial numbers that we are working with."
          # Well, did they type more after "translate" ?
          elif user_text.lower().startswith('translate'):
              # Did they enter a list of Meraki switch serial numbers after "to" ?
              serials = meraki_serials
              if re.search('to ', user_text, re.IGNORECASE):
                  if debug:
                      print(f"user_text.split('to ',1)[1] = {user_text.split('to ',1)[1]}")
                  if not user_text.split("to ",1)[1] == "":
                      serials, r = Split_check_serials(user_text,"Translate")
                      if debug:
                          print(f"serials = {serials}")
                          print(f"r = {r}")
                          print(f"r=='' = {r==''}")
                      if not r == "":
                          response.markdown = r
                          return(response)
              # Did they type "translate file <something>" ?
              if re.search('file ', user_text, re.IGNORECASE):
                  if not user_text.split("file ",1)[1] == "":
                      # They did, so translate it!
                      maybe_file = user_text.split("file ",1)[1].split()[0]
                      if debug:
                          print(f"maybe_file = {maybe_file}")
                      maybe_file, exists = File_exists(maybe_file)
                      if not exists:
                          response.markdown = "I'm sorry, but I could not find that file."
                      else:
                          response.markdown = translate_switch(incoming_msg,config=maybe_file,serials=serials)
                  else:
                      # They did not, so BUMP the user.
                      response.markdown = "I'm sorry, but I don't have a config that we are working with.  Use the **/check** command."
              # Did they type "translate host <something>" ?
              elif re.search('host ', user_text, re.IGNORECASE):
                  if debug:
                      print(f"I made it to host...")
                  if not user_text.split("host ",1)[1] == "":
                      # They did, so translate it!
                      maybe_host = user_text.split("host ",1)[1]
                      regex = re.compile(r"\s*to\s *", flags=re.I)
                      the_rest = regex.split(maybe_host)[0]
                      host_id = the_rest.strip()
                      if debug:
                          print(f"host_id = {host_id}")
                      if not Ping(host_id):
                          response.markdown = "I was unable to ping that host."
                      response.markdown = translate_switch(incoming_msg,host=host_id,serials=serials)
                      return(response)
                  else:
                      # They did not, so BUMP the user.
                      response.markdown = "I'm sorry, but I don't have a host that we are working with.  Use the **/check** command."
              else:
                  if debug:
                      print(f"len(serials) = {len(serials)}")
                  if len(serials) == 0:
                      if debug:
                          print(f"Why am I here???")
                      response.markdown = "You need to enter a list of Meraki switch serial numbers."
                  else:
                      response.markdown = translate_switch(incoming_msg,serials=serials)
        
        case "check":
            if user_text.lower() == 'check':
                if host_id == "" and config_file == "":
                    # We did not...
                    response.markdown = "I'm sorry, but I don't know what switch or filespec we are working with.  Use the **/check** command."
                else:
                    # We did, so check it!
                    response.markdown = check_switch(incoming_msg,host=host_id,config=config_file)
            elif not len(user_text.split()) >= 3:
                response.markdown = "Syntax is **check (host <_fqdn or ip address_> | file <_filespec_>)**"
            elif user_text.lower().startswith('check'):
                if re.search('file', user_text, re.IGNORECASE):
                    if not user_text.split("file ",1)[1] == "":
                        maybe_file = user_text.split("file ",1)[1].split()[0]
                        maybe_file, exists = File_exists(maybe_file)
                        if not exists:
                            response.markdown = "I'm sorry, but I could not find that file."
                        else:
                            response.markdown = check_switch(incoming_msg,config=maybe_file)
                    else:
                        response.markdown = "I'm sorry, but I don't have a config that we are working with.  Use the **/check** command."
                elif re.search('host ', user_text, re.IGNORECASE):
                    if not user_text.lower().split("host ",1)[1] == "":
                        host_id = user_text.lower().split("host ",1)[1]
                        if debug:
                            print(f"Ping({host_id}) = {Ping(host_id)}")
                        if not Ping(host_id):
                            response.markdown = "I was unable to ping that host."
                            return(response)
                        if BOT:
                            response.html = check_switch(incoming_msg,host=host_id)
                        else:
                            response.markdown = check_switch(incoming_msg,host=host_id)
                    else:
                        response.markdown = "I'm sorry, but I don't have a host that we are working with.  Use the **/check** command."
            else:
                response.markdown = "Syntax is **check (host <_fqdn or ip address_> | file <_filespec_>)**"
        
        case "register":
            # If the only thing the user typed was register...
            if user_text.lower() == 'register':
                # Check and see if we have a global stateful host to work with
                if host_id == "":
                    # We did not...
                    response.markdown = "I'm sorry, but I don't have a host that we are working with.  Use the **/check** command."
                else:
                     # We did, so translate it!
                    response.markdown = register_switch(incoming_msg,host=host_id)
            # Well, did they type more after "register" ?
            elif user_text.lower().startswith('register'):
                # Did they type "register file <something>" ?
                if re.search('host ', user_text, re.IGNORECASE):
                    if not user_text.split("host ",1)[1] == "":
                        # They did, so register it!
                        host_id = user_text.split("host ",1)[1]
                        if not Ping(host_id):
                            response.markdown = "I was unable to ping that host."
                            return(response)
                        response.markdown = register_switch(incoming_msg,host=host_id)
                    else:
                        # They did not, so BUMP the user.
                        response.markdown = "I'm sorry, but I don't have a host that we are working with.  Use the **/check** command."
                else:
                    response.markdown = "Either enter **register host _fqdn or ip address_**."
        
        case "claim":
            # If the only thing the user typed was register...
            if user_text.lower() == 'claim':
                # Check and see if we have a global stateful list of Meraki serial numbers to work with
                if len(meraki_serials) == 0 and meraki_net == "":
                    # We did not...
                    response.markdown = "I'm sorry, but I don't have a Network a list of Meraki serial numbers that we are working with."
                else:
                    if len(meraki_serials) == 0:
                        # We did not...
                        response.markdown = "I'm sorry, but I don't have a list of Meraki serial numbers that we are working with."
                    elif meraki_net == "":
                        # We did not...
                        response.markdown = "I'm sorry, but I don't have a Network that we are working with."
                    else:
                        # We did, so claim it!
                        response.markdown = claim_switch(incoming_msg,dest_net=meraki_net,serials=meraki_serials)
            
            # Well, did they type more after "claim" ?
            elif user_text.lower().startswith('claim'):
                # Did they enter a Meraki network after "to" ?
                dest_net = meraki_net
                if re.search('to ', user_text, re.IGNORECASE):
                    regex = re.compile(r"\s*to\s *", flags=re.I)
                    if not regex.split(user_text)[1] == "":
                        # They did, so let's grab and test it
                        dest_net_name = regex.split(user_text)[1]
                        user_text = re.sub('\sto\s'+dest_net_name, '', user_text, re.IGNORECASE)
                        test_net = [d['id'] for d in meraki_networks if d['name']==dest_net_name]
                        if not len(test_net) == 0:
                            dest_net = test_net[0]
                            meraki_net_name = dest_net_name
                        else:
                            response.markdown = "I'm sorry, but {} is not in your list of Meraki networks.".format(dest_net_name)
                # Set the list of serial numbers to the global stateful list in case
                # nothing was entered by the user
                serials = list()
                # Did the user enter a list of serial numbers after "Claim" ?
                if not user_text.lower() == 'claim':
                    maybe_serials, r = Split_check_serials(user_text,"Claim")
                    if debug:
                        print(f"In case: 'claim': after Split_check_serials, serials = {serials}, maybe_serials = {maybe_serials}, r = {r}")
                    # If we didn't get back some kind of error response, use the returned list
                    if r == "":
                        serials = maybe_serials
                    else:
                        response.markdown = r
                else:
                    serials = meraki_serials
                # NOW let's see if we have any serial numbers and a Network to work with...
                if len(serials)==0 and dest_net=="":
                    # No to both
                    response.markdown = "Try entering **claim _meraki serial numbers_ to _meraki Network name_**."
                else:
                    if len(serials) == 0:
                        # Just missing serial numbers
                        response.markdown = "I'm sorry, but I don't have a list of Meraki serial numbers that we are working with."
                    elif dest_net == "":
                        # Just missing a network
                        response.markdown = "I'm sorry, but I don't have a Network that we are working with."
                    else:
                        # All good, let's go claim the serial numbers to the network
                        r = claim_switch(incoming_msg,dest_net=dest_net,serials=serials)
                        response.markdown = r
        
        case "help":
            if BOT:
                # Lookup details about sender for our default response
                sender = bot.teams.people.get(incoming_msg.personId)
                response.markdown = "Well {}, here's a list of commands that I understand. ".format(sender.firstName)
                for line in bot_commands:
                    response.markdown += "\n" + line[0] + ": " + line[1]
            else:
                response.markdown = "\n\n" + tabulate(command_list,headers=["Command Format","Function"]) + "\n"
        
        case "hi" | "hello":
            if BOT:
                # Lookup details about sender for our default response
                sender = bot.teams.people.get(incoming_msg.personId)
                response.markdown = "Hi, {}! ".format(sender.firstName)
                response.markdown += "What do you want me to do today?\n"
                response.markdown += "See what I can do by asking for **help**."
        
        case _:
            if BOT:
                # Lookup details about sender for our default response
                sender = bot.teams.people.get(incoming_msg.personId)
                response.markdown = "Hello {}, I'm really just a glorified chat bot. ".format(sender.firstName)
                response.markdown += "See what I can do by asking for **help**."
            else:
                response.markdown = "Hello, I'm really just a glorified chat bot. "
                response.markdown += "See what I can do by asking for help."
    
    # Whatever just happened up above, send our response back to the user.
    return response



# Create functions that will be linked to bot commands to add capabilities
# ------------------------------------------------------------------------


# This function will check a Catalyst switch config for compatible Meraki features
#
def check_switch(incoming_msg,config="",host=""):
    """
    This function will check a Catalyst switch config for feature mapping to Meraki.
    :param incoming_msg: The incoming message object from Teams
    :param config: The incoming config filespec from the check_card Teams card
    :param host: The incoming hostname or IP address from the check_card Teams card
    :return: A text or markdown based reply
    """
    
    start_time = time.time()
    
    # Import the global stateful variables
    global config_file, host_id, times
    
    if config == "":
        
        # Since we weren't passed a config filespec, check for a hostname or IP address
        if host == "":
            return "You need to enter either a host or a filename."
        
        # We were passed a hostname or IP address...
        else:
            host_id = host
        
        # SSH to the switch with netmiko, read the config, grab the hostname,
        # write the config out to a file using the hostname as part of the filespec
        session_info = {
            'device_type': 'cisco_xe',
            'host': host_id,
            'username': ios_username,
            'password': ios_password,
            'port' : ios_port,          # optional, defaults to 22
            'secret': ios_secret,     # optional, defaults to ''
        }
        net_connect = ConnectHandler(**session_info)
        switch_name = net_connect.find_prompt()
        net_connect.enable()
        switch_name = net_connect.find_prompt()
        switch_name = switch_name[:len(switch_name) - 1]
        net_connect.send_command('term len 0')
        config = net_connect.send_command('show running-config')
        net_connect.send_command('term len 24')
        net_connect.disconnect()
        dir = os.path.join(os.getcwd(),"../files")
        file = open(os.path.join(dir,switch_name+".cfg"), "w")
        file.writelines(config)
        file.close()
        config = os.path.join(dir,switch_name+".cfg")
    
    # Update the global stateful variable for later
    config_file = config
    
    # Run the function in config_checker to get the list of features configured on
    # the switch (supported and not)
    # host_name,the_list,unsupported_features,More_info = CheckFeatures(config_file)
    host_name,the_list = CheckFeatures(config_file)
    
    # Clear some variables for the next step
    can_list = list()
    can_list_doc = list()
    not_list = list()
    not_list_doc = list()
    more = dict()
    # Go through the outcome from the read_conf functions and split the supported and
    # unsupported features as well as the additional text and links for the unsupported
    
    x = 0
    not_notes = list()
    while x < (len(the_list)):
        if not the_list[x][1] == "":
            can_list_doc.append(the_list[x])
            can_list.append([the_list[x][0],the_list[x][1],the_list[x][2]])
        else:
            not_list_doc.append(the_list[x])
            not_list.append([the_list[x][0]," "," "])
            not_notes.append([the_list[x][3],the_list[x][4]])
        x +=1
    all_list = list(list())
    all_list.extend(can_list)
    all_list.extend(not_list)
    
    if BOT: 
        tabulate.PRESERVE_WHITESPACE = True
        # Build the report.
        if debug:
            print(f"all_list = {all_list}")
        report=tabulate(all_list,colalign=["left","center","center"],headers=["Feature","Available","Translatable"])
        report_lines = report.splitlines()
        report_line_len = len(report_lines[1])
        x = 0
        bad_start = len(can_list) + 2
        while x < len(report_lines):
            results = [(m.start(), m.end()-1) for m in re.finditer(r'\S+', report_lines[x])]
            fix_line = ""
            last_word = 0
            for result in results:
                if not last_word == 0:
                    num = result[0] - last_word
                    if num > 1:
                        fix_line += "".join(['&nbsp'*num])
                    fix_line += " "
                fix_line+= report_lines[x][result[0]:result[1]+1]
                last_word = result[1]+1
            num = report_line_len - last_word
            fix_line += "".join(['&nbsp'*num])
            report_lines[x] = fix_line
            if x > 1:
                if all_list[x-2][1] in [" ",""]:
                    report_lines[x] += "&nbsp"
                if all_list[x-2][2] in [" ",""]:
                    report_lines[x] += "&nbsp"
            report_lines[x] = "<code>" + report_lines[x] + "</code>"
            x+=1
        x = 0
        while x < len(not_list):
            if not not_notes[x][0] == "":
                hotlink = '<a href =\"'+not_notes[x][1]+'" rel="nofollow\">'+not_notes[x][0]+'</a>'
                report_lines[bad_start + x]+=hotlink
            x+=1
        new_report = '<br>'.join(report_lines)
        new_report = "<p>" + new_report+ "</p>"
        new_report = '<h3>Merakicat Feature Report for ' + switch_name + '</h3><br>' + new_report
        fname = check_report_writer(switch_name,can_list_doc,not_list_doc)
        return(new_report + '''<br><br><b>Please review the results above.</b>
<br>If you wish, I can migrate the Translatable features to an existing switch in the Meraki Dashboard.  Type <b>translate</b> and a Meraki switch serial number.
<br>If you prefer, I can prepare for the switch to become a Meraki managed switch, keeping the translated config.  Just type <b>migrate [to <i>meraki network</i>]</b>.
''')
    
    # Not a BOT
    else:
        # Build the report
        if debug:
            print(f"all_list = {all_list}")
        report=tabulate(all_list,headers=["Feature","Available","Translatable","Notes","For more info, see this URL"])
        timing = ""
        if times == True:
            timing =  "\n=== That config check took %s seconds" % str(round((time.time() - start_time), 2))
        fname = check_report_writer(switch_name,an_list_doc,not_list_doc)
        return(report + "\n\nPlease review the results above, or in the file " + fname + ".\nIf you wish, I can translate or migrate the Translatable features to an existing switch in the Meraki Dashboard."+timing)


def check_report_writer(switch_name,can_list_doc,not_list_doc, pdf=False):
    document = docx.Document()
    section = document.sections[0]
    header = section.header
    header.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraphs[0].text = "Merakicat Feature Check Report for "+switch_name
    table = document.add_table(rows=1,cols=4)
    table.autofit = False
    col_count = len(table.columns)
    headers=["Feature","Available","Translatable","More Information"]
    heading_cells = table.rows[0].cells
    set_col_widths(table.rows[0])

    x = 0
    while x< col_count:
        heading_cells[x].text = headers[x]
        heading_cells[x].paragraphs[0].runs[0].font.bold = True  # Bold
        x += 1
    heading_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_repeat_table_header(table.rows[0])
    for line in can_list_doc:
        row = table.add_row()
        set_col_widths(row)
        cells = row.cells
        x = 0
        while x< col_count-1:
            cells[x].text = line[x]
            x += 1
        cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in not_list_doc:
        row = table.add_row()
        set_col_widths(row)
        cells = row.cells
        cells[2].merge(cells[3])
        x = 0
        while x< col_count-1:
            cells[x].text = line[x]
            x += 1
        p_table = cells[2].paragraphs[0]
        p_table.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hyperlink = add_hyperlink(p_table, line[3], line[4], '0000FF', False)
    dir = os.path.join(os.getcwd(),"../files")
    fname = switch_name+".docx"
    fname_pdf = switch_name+".pdf"
    document.save(os.path.join(dir,fname))
    if pdf:
        convert(os.path.join(dir,fname), (os.path.join(dir,fname_pdf)))
        os.remove(os.path.join(dir,fname))
        fname = fname_pdf
    return(fname)

#document.tables[0].rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(50, 0, 255)  # Blue Color
#document.tables[0].rows[0].cells[0].paragraphs[0].runs[0].font.blod = True  # Bold


def add_table_cell(table, row, col, text, fontSize=8, r=0, g=0, b=0, width=-1):
    cell = table.cell(row,col)
    if (width!=-1):
        cell.width = Inches(width)
    para = cell.add_paragraph(style=None)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = False
    run.font.size = Pt(fontSize)
    run.font.color.type == MSO_COLOR_TYPE.RGB
    run.font.color.rgb = RGBColor(r, g, b)


def set_col_widths(row):
    widths = (Inches(2), Inches(1), Inches(1.1), Inches(2))
    for idx, width in enumerate(widths):
        row.cells[idx].width = width


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


def add_hyperlink(paragraph, text, url, color, underline):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)
    # Join all the xml elements together  add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


#This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs
    return "Hyperlink"

def register_switch(incoming_msg,host="",called=""):
    """
    This function will register a Catalyst switch to the Meraki Dashboard.
    :param incoming_msg: The incoming message object from Teams
    :param host: The incoming hostname or IP address from the translate_card Teams card
    :param called: Indicates if the function was called from another function vs greeting
    :return: A text or markdown based reply if called from greeting (called="")
    :      : Or, status, issues & registered switch list otherwise
    """
    
    start_time = time.time()
    
    # Import the global stateful variables
    global host_id, meraki_serials, nm_list, times
    
    # Clear some variables for the next step
    switch_name = ""
    
#    if config == "":
    
    # Since we weren't passed a config filespec, check for a hostname or IP address
    if host == "":
        return("You need to enter a host FQDN or IP address.")
    else:
    # We were passed a hostname or IP address...
    # Update the global stateful variable for later
        host_id = host
    
    # SSH to the switch with netmiko, read the config, grab the hostname,
    # write the config out to a file using the hostname as part of the filespec
    status, issues, registered_switches, registered_serials, nm_list = Register(host,
        ios_username,ios_password,ios_port,ios_secret)
    if debug:
        print(f"In register_switch, status = {status}")
    if status == "successfully":
        meraki_serials = registered_serials
        if debug:
            for switch in registered_switches:
                print(f"In register_switch status = {status}")
                print(f"switch = {switch}")
                print(f"switch['Migration Status'] = {switch['migration_status']}")
    if debug:
        print(f"After registering switches, meraki_serials = {meraki_serials}")
    # Report back on what happened
    if called == "":
        if not len(registered_switches) == 0:
            vals=reduce(lambda x,y:x+y,[list(dic.values()) for dic in registered_switches])
            header = registered_switches[0].keys()
            rows = [x.values() for x in registered_switches]
            thing = tabulate(rows, header)
            timing = ""
            if times == True:
                timing =  "\n=== That registraion took %s seconds" % str(round((time.time() - start_time), 2))
            if BOT:
                payload = "```\n%s"%thing + "\n```" + timing
                return (f"We **{status}** registered **{vals.count('Registered')}** switch{'es' if (vals.count('Registered') > 1) else ''}:\n{payload}")
            else:
                payload = "\n%s"%thing + timing
                return (f"\n\nWe {status} registered {vals.count('Registered')} switch{'es' if (vals.count('Registered') > 1) else ''}:\n{payload}")
        else:
            payload = ""
            for issue in issues:
                payload += issue + "\n"
            return (f"We were unsuccessful registering {host}:\n\n{payload}" + timing)        
    else:
        return (status, issues, registered_switches)


def claim_switch(incoming_msg,dest_net=meraki_net,serials=meraki_serials,called=""):
    """
    This function will Claim newly Registered Catalyst switch in the Meraki Dashboard.
    :param incoming_msg: The incoming message object from Teams
    :param dest_net: The incoming Meraki destination Network to claim devices to
    :param serials: The incoming list of Meraki serial numbers to claim
    :param called: Indicates if the function was called from another function vs greeting
    :return: A text or markdown based reply if called from greeting (called="")
    :      : Or, status, issues, already claimed & claimed switch lists otherwise
    """
    
    start_time = time.time()
    
    global host_id, meraki_net, meraki_serials, meraki_net_name, times
    issues = ""
    if debug:
        print(f"At start of claim_switch, serials = {serials}")
    claimed_switches = serials
    ac_switches = list()
    bad_switches = list()
    
    if dest_net == "":
        return (f"claim_switch was called with no dest_net!")
    if debug:
        print(f"len([d['name'] for d in meraki_networks if d['id']=={dest_net}]) = {len([d['name'] for d in meraki_networks if d['id']==dest_net])}")
    if (not dest_net == meraki_net) and (not len([d['name'] for d in meraki_networks if d['id']==dest_net]) == 1):
        return (f"claim_switch was called with a dest_net that doesn't match any of your Meraki Network IDs!")
    if len(serials) == 0:
        return (f"claim_switch was called with no serials!")
    
    issues,bad_switches,ac_switches,claimed_switches = Claim(dashboard,dest_net,serials)
    
    # If the claim went fine, update the global stateful variables for later
    if len(bad_switches) == 0:
        meraki_net = dest_net
        meraki_serials = serials
        test_net = [d['name'] for d in meraki_networks if d['id']==meraki_net]
        if not len(test_net) == 0:
            meraki_net_name = test_net[0]
    
    # Report back on what happened
    
    # If we were not called from another function,
    # if there were no Bad_switches, return a nice message
    # otherwise return the list of issues
    if called == "":
        if len(bad_switches) == 0:
            r = "I was able to claim those switches to your Network.\n"
            if times == True:
                r += "=== That claiming process took %s seconds" % str(round((time.time() - start_time), 2))
            return (r)
        else:
            return(issues)
    # If we WERE called from another function,
    # return the list of issues, lists of claimed & already claimed switches
    # and a status of "OK" if no bad switches, or a status of "Issues"
    else:
        status = "Ok" if len(bad_switches) == 0 else "Issues"
        if debug:
            print(f"At the end of our claim_switch call:")
            print(f"status = {status}")
            print(f"issues = {issues}")
            print(f"ac_switches = {ac_switches}")
            print(f"claimed_switches = {claimed_switches}")
        return (status, issues, ac_switches, claimed_switches)


def translate_switch(incoming_msg,config=config_file,host=host_id,serials=meraki_serials,verb="translate"):
    """
    This function will translate a Catalyst switch stack config to features in an existing
    set of Meraki switches.
    :param incoming_msg: The incoming message object from Teams
    :param config: The incoming config filespec from the translate_card Teams card
    :param host: The incoming hostname or IP address from the translate_card Teams card
    :param serials: The incoming list of Meraki serials from the translate_card Teams card
    :param verb: Either 'translate' or 'migrate' depending on how the function is called
    :return: A text or markdown based reply
    """
    start_time = time.time()
      
    # Import the global stateful variables
    global config_file, host_id, meraki_org, meraki_serials, meraki_urls, nm_list, times
        
    if debug:
        print(f"In translate, config_file = {config_file}")
    # Clear some variables for the next step
    switch_name = ""
    # Check whether or not we were passed a list of up to 8 Meraki serial numbers
    if len(serials) > 8:
        return "A switch stack can contain a maximum of 8 switches."
    # Update the global stateful variable for later
    meraki_serials = serials
    if debug:
        print(f"meraki_serials = {meraki_serials}")
    # Check whether or not we were passed a config filespec
    
    switch_name = ""
    if not config == "":
        ext = os.path.splitext(config)[1]
        switch_name = os.path.split(config)[1].replace(ext,"")
    elif not config_file == "":
        ext = os.path.splitext(config_file)[1]
        switch_name = os.path.split(config_file)[1].replace(ext,"")
    if config == "" and config_file == "":
        if host == "" and host_id == "":
            return "You need to enter either a host or a config filespec."
        else:
            if host == "":
                host = host_id
            host_id = host
            # SSH to the switch with netmiko, read the config, grab the switch name,
            # write the config out to a file using the switch name as part of the filespec
            if debug:
                print(f"meraki_serials = {meraki_serials}")
            session_info = {
                'device_type': 'cisco_xe',
                'host': host_id,
                'username': ios_username,
                'password': ios_password,
                'port' : ios_port,          # optional, defaults to 22
                'secret': ios_secret,     # optional, defaults to ''
            }
            if debug:
                print(f"session_info = {session_info}")
            net_connect = ConnectHandler(**session_info)
            switch_name = net_connect.find_prompt()
            net_connect.enable()
            switch_name = net_connect.find_prompt()
            switch_name = switch_name[:len(switch_name) - 1]
            net_connect.send_command('term len 0')
            config = net_connect.send_command('show running-config')
            net_connect.send_command('term len 24')
            net_connect.disconnect()
            dir = os.path.join(os.getcwd(),"../files")
            file = open(os.path.join(dir,switch_name+".cfg"), "w")
            file.writelines(config)
            file.close()
            config = os.path.join(dir,switch_name+".cfg")
    else:
        if config == "":
            config = config_file
                
    # Update the global stateful variable for later
    config_file = config
    
    eval_start_time = time.time()
    
    # Evaluate the Catalyst config and break it into lists we can work with
    Uplink_list, Downlink_list, Other_list, port_dict, switch_dict = Evaluate(config_file)
    
    ## Creating a list of the downlink port configurations to push to Meraki
    ToBeConfigured = {}
    z = 0
    while z < len(Downlink_list):
        interface = Downlink_list[z]
        ToBeConfigured[interface] = port_dict[interface]
        z +=1

    ##
    ## Start the meraki config migration after confirmation from the user
    ##
    blurb = "Evaluated the switch config."
    if times == True:
        blurb += "\n--- That took %s seconds" % str(round((time.time() - start_time), 2))
    blurb += "\n\nPushing the translated items to the Dashboard in a large batch.\nThis will take a while, but I'll message you when I'm done..."
    if BOT:
        c = create_message(incoming_msg.roomId, blurb)
    else:
        print(blurb)
    
    port_cfg_start_time = time.time()
    
    configured_ports,unconfigured_ports,port_dict,meraki_urls,meraki_net = Meraki_config_down(dashboard,meraki_org,switch_name,meraki_serials,port_dict,Downlink_list,Other_list,switch_dict)
        
    if debug:
        print(f"configured_ports = {configured_ports}")
        print(f"unconfigured_ports = {unconfigured_ports}")
    '''
    ## Creating a list of the uplink port configurations to push to Meraki
    ToBeConfigured = {}
    z = 0
    while z < len(Uplink_list):
        interface = Uplink_list[z]
        ToBeConfigured[interface] = port_dict[interface]
        z +=1
    
    ## Creating a list of the uplink port configurations to push to Meraki
    ChannelsToBeConfigured = {}
    z = 0
    while z < len(Other_list):
        interface = Other_list[z]
        ChannelsToBeConfigured[interface] = port_dict[interface]
        z +=1
    
    ##
    ## Start the meraki uplink config migration after confirmation from the user
    ##
    if BOT:
        c = create_message(incoming_msg.roomId, "Pushing the translated uplinks and Port-channels to the Dashboard, port by port.  This will take a while, but I'll message you when I'm done...")
    else:
        print("Pushing the translated uplinks and Port-channels to the Dashboard, port by port.  This will take a while, but I'll let you know when I'm done...")
    configured_up_ports,unconfigured_up_ports,configured_channels,unconfigured_channels = Meraki_config_up(dashboard,meraki_org,meraki_net,meraki_serials,ToBeConfigured,Uplink_list,nm_list,ChannelsToBeConfigured,Other_list)
    
    if debug:
        print(f"configured_up_ports = {configured_up_ports}")
        print(f"unconfigured_up_ports = {unconfigured_up_ports}")
    '''
    dir = os.path.join(os.getcwd(),"../files")
    with open(os.path.join(dir,switch_name+".pd"), 'w') as file:
        file.write(json.dumps(port_dict)) # use `json.loads` to do the reverse
        file.close()

    switch = 0
    r = ""
    if debug:
        print(f"meraki_serials = {meraki_serials}")
    while switch <= len(meraki_serials)-1:
        if len(meraki_serials) == 1:
            r += "\nFor the switch ["+meraki_serials[switch]+"]("+meraki_urls[switch]+"):\n"
        else:
            r += "\nFor switch "+str(switch+1)+" ["+meraki_serials[switch]+"]("+meraki_urls[switch]+"):\n"
        if len(configured_ports[switch]) > 0:
            if BOT:
                r += "We were able to **successfully** " + verb + " ports: "
            else:
                r += "We were able to successfully " + verb + " ports: "
            c_port = 0
            while c_port <= len(configured_ports[switch])-2:
                r += configured_ports[switch][c_port] + ", "
                c_port+=1
            r += configured_ports[switch][c_port] + "\n\n"
        if len(unconfigured_ports[switch]) > 0:
            if BOT:
                r += "We were **unable** to " + verb + " ports: "
            else:
                r += "We were unable to " + verb + " ports: "
            u_port = 0
            while u_port <= len(unconfigured_ports[switch])-2:
                r += unconfigured_ports[switch][u_port] + ", "
                u_port+=1
            r += unconfigured_ports[switch][u_port] + "\n\n"
        switch+=1
    if verb == "translate" and times:
        r += "\n--- Pushing to Dashboard took %s seconds" % str(round((time.time() - port_cfg_start_time), 2))
        r += "\n=== That entire translation took %s seconds" % str(round((time.time() - start_time), 2))
    return (r)


def migrate_switch(incoming_msg,host=host_id,dest_net=meraki_net):
    """
    This function will register a Catalyst switch stack to the Meraki Dashboard, claim
    the stack to a Meraki Network, then translate the switch stack config to Meraki.
    Once finished, the user can edit the Meraki stack config before manually initiating
    migration to Cloud Management via "service meraki start" CLI command on the stack.
    :param incoming_msg: The incoming message object from Teams
    :param host: The incoming hostname or IP address
    :param dest_net: The incoming Meraki destination Network to claim devices to
    :return: A text or markdown based reply
    """
    
    start_time = time.time()
    
    # Import the global stateful variables
    global config_file, host_id, nm_list, times
    global meraki_net, meraki_net_name, meraki_serials, meraki_urls
    
    # Clear some variables for the next step
    switch_name = ""
    status = ""
    issues = ""
    ac_switches = list()
    claimed_switches = list()
    
    if debug:
        print(f"At the start of migrate_switch:")
        print(f"host = {host}")
        print(f"dest_net = {dest_net}")
    # Were we passed a hostname or IP address?
    if host == "":
        return "You need to provide a host."
    else:
    # We were passed a hostname or IP address...
        host_id = host
    
    # Were we passed a Meraki Network?
    if dest_net == "":
        return "You need to provide a Meraki network."
    else:
        if debug:
            print(f"")
    # We were passed a hostname or IP address...
        # Is it in the list of the user's Meraki networks?
        if debug:
            print(f"meraki_networks = {meraki_networks}")
        if (not dest_net == meraki_net):
            if debug:
                print(f"In migrate_switch, {dest_net} != {meraki_net}")
            if (not len([d['name'] for d in meraki_networks if d['id']==dest_net]) == 1):
                if debug:
                    print(f"In migrate_switch, {dest_net} != {[d['id'] for d in meraki_networks]}")
                return "You need to provide a Meraki network."
        # It was in the list of the user's Meraki networks, so save it
        meraki_net = dest_net
        test_net = [d['name'] for d in meraki_networks if d['id']==meraki_net]
        if not len(test_net) == 0:
            meraki_net_name = test_net[0]

    # SSH to the switch with netmiko, read the config, grab the switch name,
    # write the config out to a file using the switch name as part of the filespec
    session_info = {
        'device_type': 'cisco_xe',
        'host': host_id,
        'username': ios_username,
        'password': ios_password,
        'port' : ios_port,          # optional, defaults to 22
        'secret': ios_secret,     # optional, defaults to ''
    }
    net_connect = ConnectHandler(**session_info)
    switch_name = net_connect.find_prompt()
    net_connect.enable()
    switch_name = net_connect.find_prompt()
    switch_name = switch_name[:len(switch_name) - 1]
    net_connect.send_command('term len 0')
    config = net_connect.send_command('show running-config')
    net_connect.send_command('term len 24')
    net_connect.disconnect()
    dir = os.path.join(os.getcwd(),"../files")
    file = open(os.path.join(dir,switch_name+".cfg"), "w")
    file.writelines(config)
    file.close()
    config = os.path.join(dir,switch_name+".cfg")
    blurb = "Logged in to " + host_id + ", grabbed a copy of the running config and saved it as " + switch_name+".cfg."
    if times == True:
        blurb += "\n--- That took %s seconds" % str(round((time.time() - start_time), 2))
    if BOT:
        c = create_message(incoming_msg.roomId, blurb)
    else:
        print(blurb)
    
    # Update the global stateful variable for later
    config_file = config
    
    # Register the switch stack to the Meraki dashboard
    if debug:
        print(f"in migrate before register_switch, meraki_serials = {meraki_serials}")    
    
    register_start_time = time.time()
    
    status, issues, registered_switches = register_switch(incoming_msg,host=host_id,called='yes')

    if debug:
        print(f"in migrate after register_switch, meraki_serials = {meraki_serials}")    
    
    # If we were not fully successful, just return with the report
    if not status == "successfully":
        vals=reduce(lambda x,y:x+y,[list(dic.values()) for dic in registered_switches])
        header = registered_switches[0].keys()
        rows = [x.values() for x in registered_switches]
        thing = tabulate(rows, header)
        payload = "```\n%s"%thing
        return (f"We **{status}** registered **{vals.count('Registered')}** switch{'es' if (vals.count('Registered') > 1) else ''}:\n{payload}")
    string_serials = ', '.join(meraki_serials)
    blurb = "Registered " + host_id + " to Dashboard as " + string_serials + ".\n"
    if times == True:
        blurb += "--- That took %s seconds" % str(round((time.time() - register_start_time), 2))
    if BOT:
        c = create_message(incoming_msg.roomId, blurb)
    else:
        print(blurb)
    if debug:
        print(f"in migrate before claim_switch, meraki_serials = {meraki_serials}")    
    # Claim the switch stack to a Network in the Meraki dashboard
    claim_start_time = time.time()
    status, issues, ac_switches, claimed_switches = claim_switch(incoming_msg,dest_net=meraki_net,serials=meraki_serials,called='yes')
    if debug:
        print(f"in migrate after claim_switch, meraki_serials = {meraki_serials}")    

    # If the attempt to claim the switch stack had issues, return them
    if not status == "Ok":
        return(issues)
    blurb = "Claimed " + string_serials + " to Meraki network " + meraki_net_name + ".\n"
    if times == True:
        blurb += "--- That took %s seconds" % str(round((time.time() - claim_start_time), 2))
    if BOT:
        c = create_message(incoming_msg.roomId, blurb)
    else:
        print(blurb)
    if debug:
        print(f"in migrate before translate, meraki_serials = {meraki_serials}")
    # Translate the switch stack to the Meraki switches we just claimed
    translate_start_time = time.time()
    r = "\n\n" + translate_switch(incoming_msg,config="",host=host_id,serials=meraki_serials,verb="migrate")
    blurb = "\nTranslated " + switch_name+".cfg to Meraki switches " + string_serials + ".\n"
    if times == True:
        blurb += "--- That took %s seconds" % str(round((time.time() - translate_start_time), 2))
        blurb += "\n=== For a total time of  %s seconds" % str(round((time.time() - start_time), 2))
    if BOT:
        c = create_message(incoming_msg.roomId, blurb)
    else:
        print(blurb)
    r += "\nPlease review the configuration in Dashboard and add/modify what you want prior to converting the switch to Meraki Cloud Management."
    if BOT:
        r += "\n**Converting the switch will _remove all configuration_, so you do so at your own risk!**"
    else:
        r += "\nCONVERTING THE SWITCH WILL REMOVE ALL CONFIGURATION, SO YOU DO SO AT YOUR OWN RISK!"
    r += "\n To convert the switch, enter the following:\n"
    if BOT:
        r += "```\nenable\nservice meraki start"
    else:
        r += "\n    enable\n    service meraki start"
    return (r)


def create_message(rid, msgtxt):
    headers = {
        "content-type": "application/json; charset=utf-8",
        "authorization": "Bearer " + teams_token,
    }
    
    url = "https://api.ciscospark.com/v1/messages"
    data = {"roomId": rid, "markdown": msgtxt}
    response = requests.post(url, json=data, headers=headers)
    return response.json()


# Temporary function to send a message with a card attachment (not yet
# supported by webexteamssdk, but there are open PRs to add this
# functionality)
def create_message_with_attachment(rid, msgtxt, attachment):
    headers = {
        "content-type": "application/json; charset=utf-8",
        "authorization": "Bearer " + teams_token,
    }
    
    url = "https://api.ciscospark.com/v1/messages"
    data = {"roomId": rid, "attachments": [attachment], "markdown": msgtxt}
    response = requests.post(url, json=data, headers=headers)
    return response.json()


## If we are in BOT mode, set up some bot stuff
if BOT:
    
    # Set the bot greeting.
    bot.set_greeting(greeting)
    
    # Add new commands to the bot.
    bot_commands = list(list())
    bot_commands.extend([["* **help**", "Get help."],
    ["* **check [host _FQDN or IP address_ | file _filespec_] [with timing]**", "Check a Catalyst switch config for both translatable and possible Meraki features"],
    ["* **register [host _FQDN or IP address_] [with timing]**", "Register a Catalyst switch to the Meraki Dashboard"],
    ["* **claim [_Meraki serial numbers_] [to _Meraki network name_] [with timing]**", "Claim Catalyst switches to a Meraki Network"],
    ["* **translate [host _FQDN or IP address_ | file _filespec_] [to _Meraki serial numbers_] [with timing]**", "Translate a Catalyst switch config from a file or host to claimed Meraki serial numbers"],
    ["* **migrate [host _FQDN or IP address_] [to _Meraki network name_] [with timing]**", "Migrate a Catalyst switch to a Meraki switch - register, claim & translate"]])
    bot.add_command("help", "This list of commands", greeting)
    bot.add_command("check [host _FQDN or IP address_ | file _filespec_] [with timing]", "Check a Catalyst switch config for both translatable and possible Meraki features", greeting)
    bot.add_command("register [host _FQDN or IP address_] [with timing]", "Register a Catalyst switch to the Meraki Dashboard", greeting)
    bot.add_command("claim [_Meraki serial numbers_] [to _Meraki network name_] [with timing]", "Claim Catalyst switches to a Meraki Network", greeting)
    bot.add_command("translate [host _FQDN or IP address_ | file _filespec_] [to _Meraki serial numbers_] [with timing]", "Translate a Catalyst switch config from a file or host to claimed Meraki serial numbers", greeting)
    bot.add_command("migrate [host _FQDN or IP address_] [to _Meraki network name_] [with timing]", "Migrate a Catalyst switch to a Meraki switch - register, claim & translate", greeting)
    
    # Every bot includes a default "/echo" command.  You can remove it, or any
    # other command with the remove_command(command) method.
    bot.remove_command("/echo")
else:
    command_list = list(list())
    command_list.extend([
      ["help", "This list of commands"],
      ["check host <FQDN or IP address> | file <filespec> [with timing]", "Check a Catalyst switch config for both translatable and possible Meraki features"],
      ["register host <FQDN or IP address> [with timing]", "Register a Catalyst switch to the Meraki Dashboard"],
      ["claim <Meraki serial numbers> to <Meraki network name> [with timing]", "Claim Catalyst switches to a Meraki Network"],
      ["translate host <FQDN or IP address> | file <filespec> to <Meraki serial numbers> [with timing]", "Translate a Catalyst switch config from a file or host to claimed Meraki serial numbers"],
      ["migrate host <FQDN or IP address> to <Meraki network name> [with timing]", "Migrate a Catalyst switch to a Meraki switch - register, claim & translate"]
    ])


## BOT or not?
if __name__ == "__main__":
    if BOT:
        # Run Bot
        bot.run(host="0.0.0.0", port=5000)
    else:
        if debug:
            print(f"The number of command line args I see is {len(sys.argv)-1}")
        args = sys.argv
        del args[0]
        if debug:
            print(f"The args are: {str(args)}")
        text = " ".join(args)
        if debug:
            print(f"The user input was: {text}")
        command_line_msg.text = text
        command_line_msg.personId = teams_emails
        if debug:
            print(f"command_line_msg = {command_line_msg}")
        print(greeting(command_line_msg).markdown)